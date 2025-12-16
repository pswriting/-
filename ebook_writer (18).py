import streamlit as st
import google.generativeai as genai
import re
import json
import io
import os
from datetime import datetime
from pathlib import Path

# ==========================================
# API 키 저장/불러오기 (로컬 파일)
# ==========================================
def get_config_path():
    """설정 파일 경로 반환"""
    home = Path.home()
    return home / ".ebook_app_config.json"

def load_saved_api_key():
    """저장된 API 키 불러오기"""
    config_path = get_config_path()
    try:
        if config_path.exists():
            with open(config_path, 'r') as f:
                config = json.load(f)
                return config.get('api_key', '')
    except Exception:
        pass
    return ''

def save_api_key(api_key):
    """API 키 저장"""
    config_path = get_config_path()
    try:
        config = {}
        if config_path.exists():
            with open(config_path, 'r') as f:
                config = json.load(f)
        config['api_key'] = api_key
        with open(config_path, 'w') as f:
            json.dump(config, f)
        return True
    except Exception:
        return False

# --- 페이지 설정 ---
st.set_page_config(
    page_title="전자책 작성 프로그램", 
    layout="wide", 
    page_icon="◆"
)

# --- 지구인사이트 스타일 CSS ---
st.markdown("""
<style>
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
    
    * { 
        font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, system-ui, sans-serif; 
    }
    
    /* 기본 요소 숨김 */
    .stDeployButton {display:none;} 
    footer {visibility: hidden;}
    #MainMenu {visibility: hidden;}
    
    /* 사이드바 토글 버튼 표시 */
    [data-testid="collapsedControl"] {
        display: flex !important;
        visibility: visible !important;
    }
    
    /* 메인 배경 - 순수 화이트 */
    .stApp {
        background: #ffffff;
    }
    
    /* 메인 영역 */
    .main .block-container {
        background: #ffffff;
        padding: 2rem 3rem;
        max-width: 1200px;
    }
    
    /* 사이드바 */
    [data-testid="stSidebar"] {
        background: #ffffff;
        border-right: 1px solid #eeeeee;
    }
    
    [data-testid="stSidebar"] * {
        color: #222222 !important;
    }
    
    [data-testid="stSidebar"] .stProgress > div > div > div > div {
        background: #222222;
        border-radius: 10px;
    }
    
    /* 모든 텍스트 - 진한 검정 */
    .stMarkdown, .stText, p, span, label, .stMarkdown p {
        color: #222222 !important;
        line-height: 1.7;
    }
    
    /* 제목 스타일링 */
    h1 { 
        color: #111111 !important; 
        font-weight: 700 !important; 
        font-size: 2rem !important;
        letter-spacing: -0.5px;
        margin-bottom: 1rem !important;
    }
    
    h2 { 
        color: #111111 !important; 
        font-weight: 700 !important;
        font-size: 1.4rem !important;
        margin-top: 2rem !important;
        margin-bottom: 1rem !important;
    }
    
    h3 { 
        color: #222222 !important; 
        font-weight: 600 !important;
        font-size: 1.1rem !important;
        margin-bottom: 0.8rem !important;
    }
    
    /* 탭 스타일 - 미니멀 라인 */
    .stTabs [data-baseweb="tab-list"] {
        background: transparent;
        gap: 0;
        border-bottom: 2px solid #eeeeee;
        padding: 0;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: transparent;
        color: #888888 !important;
        border-radius: 0;
        font-weight: 500;
        padding: 16px 24px;
        font-size: 15px;
        border-bottom: 2px solid transparent;
        margin-bottom: -2px;
        transition: all 0.2s;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        color: #222222 !important;
    }
    
    .stTabs [aria-selected="true"] {
        background: transparent !important;
        color: #111111 !important;
        font-weight: 700 !important;
        border-bottom: 2px solid #111111 !important;
    }
    
    /* 버튼 스타일 - 검정 배경 + 흰색 글씨 */
    .stButton > button { 
        width: 100%; 
        border-radius: 30px; 
        font-weight: 600; 
        background: #111111 !important;
        color: #ffffff !important;
        border: none !important;
        padding: 14px 32px;
        font-size: 15px;
        transition: all 0.2s;
        box-shadow: none;
    }
    
    .stButton > button:hover { 
        background: #333333 !important;
        color: #ffffff !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        transform: translateY(-1px);
    }
    
    .stButton > button:active {
        transform: translateY(0);
    }
    
    /* 버튼 내부 텍스트 강제 흰색 */
    .stButton > button p,
    .stButton > button span,
    .stButton > button div,
    .stButton > button * {
        color: #ffffff !important;
    }
    
    /* 다운로드 버튼 */
    .stDownloadButton > button {
        background: #2d5a27 !important;
        color: #ffffff !important;
        border-radius: 30px;
    }
    
    .stDownloadButton > button:hover {
        background: #3d7a37 !important;
    }
    
    .stDownloadButton > button p,
    .stDownloadButton > button span,
    .stDownloadButton > button * {
        color: #ffffff !important;
    }
    
    /* 입력 필드 */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        background: #ffffff !important;
        border: 1px solid #dddddd !important;
        border-radius: 8px !important;
        color: #222222 !important;
        padding: 14px 16px !important;
        font-size: 15px !important;
    }
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: #111111 !important;
        box-shadow: none !important;
    }
    
    .stTextInput > div > div > input::placeholder,
    .stTextArea > div > div > textarea::placeholder {
        color: #aaaaaa !important;
    }
    
    /* 셀렉트박스 */
    .stSelectbox > div > div {
        background: #ffffff !important;
        border: 1px solid #dddddd !important;
        border-radius: 8px !important;
    }
    
    .stSelectbox > div > div > div {
        color: #222222 !important;
    }
    
    /* 메트릭 */
    [data-testid="stMetricValue"] {
        color: #111111 !important;
        font-size: 2rem !important;
        font-weight: 700 !important;
    }
    
    [data-testid="stMetricLabel"] {
        color: #666666 !important;
    }
    
    /* 알림 메시지 */
    .stSuccess {
        background: #f0f9f0 !important;
        border: 1px solid #c8e6c9 !important;
        border-radius: 8px !important;
    }
    .stSuccess p { color: #2e7d32 !important; }
    
    .stWarning {
        background: #fff8e1 !important;
        border: 1px solid #ffecb3 !important;
        border-radius: 8px !important;
    }
    .stWarning p { color: #f57c00 !important; }
    
    .stError {
        background: #ffebee !important;
        border: 1px solid #ffcdd2 !important;
        border-radius: 8px !important;
    }
    .stError p { color: #c62828 !important; }
    
    .stInfo {
        background: #e3f2fd !important;
        border: 1px solid #bbdefb !important;
        border-radius: 8px !important;
    }
    .stInfo p { color: #1565c0 !important; }
    
    /* 구분선 */
    hr {
        border: none !important;
        border-top: 1px solid #eeeeee !important;
        margin: 2rem 0 !important;
    }
    
    /* 프로그레스 바 */
    .stProgress > div > div > div > div {
        background: #222222;
        border-radius: 10px;
    }
    
    /* ===== 커스텀 컴포넌트 ===== */
    
    /* 로그인 화면 */
    .login-container {
        max-width: 400px;
        margin: 100px auto;
        padding: 40px;
        background: #ffffff;
        border: 1px solid #eeeeee;
        border-radius: 20px;
        text-align: center;
    }
    
    .login-title {
        font-size: 28px;
        font-weight: 700;
        color: #111111;
        margin-bottom: 8px;
    }
    
    .login-subtitle {
        font-size: 15px;
        color: #888888;
        margin-bottom: 30px;
    }
    
    /* 히어로 섹션 */
    .hero-section {
        text-align: center;
        padding: 60px 20px;
        margin-bottom: 40px;
    }
    
    .hero-label {
        font-size: 13px;
        font-weight: 600;
        color: #666666;
        letter-spacing: 3px;
        margin-bottom: 16px;
        text-transform: uppercase;
    }
    
    .hero-title {
        font-size: 42px;
        font-weight: 800;
        color: #111111;
        margin-bottom: 16px;
        letter-spacing: -1px;
        line-height: 1.2;
    }
    
    .hero-subtitle {
        font-size: 18px;
        color: #666666;
        font-weight: 400;
    }
    
    /* 섹션 라벨 */
    .section-label {
        font-size: 12px;
        font-weight: 600;
        color: #888888;
        letter-spacing: 2px;
        margin-bottom: 8px;
        text-transform: uppercase;
    }
    
    /* 점수 카드 */
    .score-card {
        background: #f8f8f8;
        border-radius: 20px;
        padding: 50px 40px;
        text-align: center;
    }
    
    .score-number {
        font-size: 80px;
        font-weight: 800;
        color: #111111;
        line-height: 1;
        margin-bottom: 8px;
    }
    
    .score-label {
        color: #888888;
        font-size: 14px;
        font-weight: 500;
    }
    
    /* 상태 배지 */
    .status-badge {
        display: inline-block;
        padding: 8px 20px;
        border-radius: 20px;
        font-weight: 600;
        font-size: 13px;
        margin-top: 20px;
    }
    
    .status-excellent {
        background: #111111;
        color: #ffffff;
    }
    
    .status-good {
        background: #f0f0f0;
        color: #333333;
    }
    
    .status-warning {
        background: #fff3e0;
        color: #e65100;
    }
    
    /* 정보 카드 */
    .info-card {
        background: #f8f8f8;
        border-radius: 16px;
        padding: 24px;
        margin: 16px 0;
    }
    
    .info-card-title {
        font-size: 12px;
        font-weight: 700;
        color: #888888;
        letter-spacing: 1px;
        margin-bottom: 12px;
        text-transform: uppercase;
    }
    
    .info-card p {
        color: #333333 !important;
        font-size: 15px;
        line-height: 1.8;
        margin: 8px 0;
    }
    
    /* 제목 카드 */
    .title-card {
        background: #ffffff;
        border: 1px solid #eeeeee;
        border-radius: 16px;
        padding: 24px;
        margin: 12px 0;
        transition: all 0.2s;
    }
    
    .title-card:hover {
        border-color: #cccccc;
        box-shadow: 0 4px 20px rgba(0,0,0,0.06);
    }
    
    .title-card .card-number {
        font-size: 12px;
        font-weight: 600;
        color: #aaaaaa;
        margin-bottom: 8px;
    }
    
    .title-card .main-title {
        color: #111111;
        font-size: 18px;
        font-weight: 700;
        margin-bottom: 6px;
    }
    
    .title-card .sub-title {
        color: #666666;
        font-size: 14px;
        margin-bottom: 16px;
    }
    
    .title-card .reason {
        color: #444444;
        font-size: 14px;
        padding: 14px 16px;
        background: #f8f8f8;
        border-radius: 10px;
        line-height: 1.6;
    }
    
    /* 점수 아이템 */
    .score-item {
        background: #ffffff;
        border: 1px solid #eeeeee;
        border-radius: 12px;
        padding: 16px 20px;
        margin: 10px 0;
        display: flex;
        justify-content: space-between;
        align-items: center;
    }
    
    .score-item-label {
        color: #333333;
        font-weight: 500;
        font-size: 15px;
    }
    
    .score-item-value {
        color: #111111;
        font-weight: 700;
        font-size: 20px;
    }
    
    .score-item-reason {
        color: #666666;
        font-size: 14px;
        margin-top: 4px;
        line-height: 1.5;
    }
    
    /* 요약 박스 */
    .summary-box {
        background: #f8f8f8;
        border-radius: 12px;
        padding: 20px;
        margin-top: 20px;
    }
    
    .summary-box p {
        color: #333333 !important;
        font-size: 15px;
        line-height: 1.7;
    }
    
    /* 푸터 */
    .premium-footer {
        text-align: center;
        padding: 40px 20px;
        margin-top: 60px;
        border-top: 1px solid #eeeeee;
    }
    
    .premium-footer-text {
        color: #888888;
        font-size: 14px;
    }
    
    .premium-footer-author {
        color: #222222;
        font-weight: 600;
    }
    
    /* 빈 상태 */
    .empty-state {
        text-align: center;
        padding: 60px 20px;
        background: #f8f8f8;
        border-radius: 16px;
    }
    
    .empty-state p {
        color: #888888 !important;
    }
    
    /* 퀵 액션 박스 */
    .quick-action-box {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        border: 1px dashed #dee2e6;
        border-radius: 16px;
        padding: 24px;
        margin: 16px 0;
        text-align: center;
    }
    
    .quick-action-box p {
        color: #495057 !important;
        font-size: 14px;
        margin-bottom: 12px;
    }
    
    /* 모드 선택 라디오 버튼 스타일 */
    .stRadio > div {
        display: flex;
        gap: 16px;
    }
    
    .stRadio > div > label {
        background: #f8f8f8;
        padding: 12px 20px;
        border-radius: 8px;
        cursor: pointer;
        transition: all 0.2s;
    }
    
    .stRadio > div > label:hover {
        background: #eeeeee;
    }
    
    /* 소제목 카드 */
    .subtopic-card {
        background: #ffffff;
        border: 1px solid #e0e0e0;
        border-radius: 12px;
        padding: 16px;
        margin: 8px 0;
    }
    
    /* 추가/삭제 버튼 작게 */
    .small-btn {
        font-size: 12px !important;
        padding: 6px 12px !important;
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 비밀번호 설정 (여기서 변경 가능)
# ==========================================
CORRECT_PASSWORD = "cashmaker2024"
# ==========================================

# --- 비밀번호 확인 ---
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

if not st.session_state['authenticated']:
    st.markdown("""
    <div class="login-container">
        <div class="login-title">CASHMAKER</div>
        <div class="login-subtitle">전자책 작성 프로그램</div>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        password_input = st.text_input("비밀번호를 입력하세요", type="password", placeholder="비밀번호")
        
        if st.button("입장하기"):
            if password_input == CORRECT_PASSWORD:
                st.session_state['authenticated'] = True
                st.rerun()
            else:
                st.error("비밀번호가 틀렸습니다")
    
    st.stop()

# --- 세션 초기화 ---
default_states = {
    'topic': '',
    'target_persona': '',
    'pain_points': '',
    'one_line_concept': '',
    'outline': [],
    'chapters': {},
    'current_step': 1,
    'market_analysis': '',
    'book_title': '',
    'subtitle': '',
    'topic_score': None,
    'topic_verdict': None,
    'score_details': None,
    'generated_titles': None,
    'outline_mode': 'ai',
}

for key, value in default_states.items():
    if key not in st.session_state:
        st.session_state[key] = value

# --- 사이드바 ---
with st.sidebar:
    st.markdown("### Progress")
    
    progress_items = [
        bool(st.session_state['topic']),
        bool(st.session_state['target_persona']),
        bool(st.session_state['outline']),
        len(st.session_state['chapters']) > 0,
        any(ch.get('content') for ch in st.session_state['chapters'].values()) if st.session_state['chapters'] else False
    ]
    progress = sum(progress_items) / len(progress_items) * 100
    
    st.progress(progress / 100)
    st.caption(f"{progress:.0f}% 완료")
    
    st.markdown("---")
    st.markdown("### Info")
    if st.session_state['topic']:
        st.caption(f"주제: {st.session_state['topic']}")
    if st.session_state['book_title']:
        st.caption(f"제목: {st.session_state['book_title']}")
    if st.session_state['outline']:
        st.caption(f"목차: {len(st.session_state['outline'])}개")
    
    completed_chapters = sum(1 for ch in st.session_state['chapters'].values() if ch.get('content'))
    if completed_chapters:
        st.caption(f"완성: {completed_chapters}개")
    
    st.markdown("---")
    st.markdown("### 💾 저장/불러오기")
    
    save_data = {
        'topic': st.session_state.get('topic', ''),
        'target_persona': st.session_state.get('target_persona', ''),
        'pain_points': st.session_state.get('pain_points', ''),
        'one_line_concept': st.session_state.get('one_line_concept', ''),
        'outline': st.session_state.get('outline', []),
        'chapters': st.session_state.get('chapters', {}),
        'book_title': st.session_state.get('book_title', ''),
        'subtitle': st.session_state.get('subtitle', ''),
        'market_analysis': st.session_state.get('market_analysis', ''),
        'topic_score': st.session_state.get('topic_score'),
        'topic_verdict': st.session_state.get('topic_verdict'),
        'score_details': st.session_state.get('score_details'),
        'generated_titles': st.session_state.get('generated_titles'),
    }
    
    save_json = json.dumps(save_data, ensure_ascii=False, indent=2)
    file_name = st.session_state.get('book_title', '전자책') or '전자책'
    file_name = re.sub(r'[^\w\s가-힣-]', '', file_name)[:20]
    
    st.download_button(
        "📥 작업 저장하기",
        save_json,
        file_name=f"{file_name}_{datetime.now().strftime('%m%d_%H%M')}.json",
        mime="application/json",
        use_container_width=True
    )
    
    uploaded_file = st.file_uploader(
        "📤 작업 불러오기",
        type=['json'],
        label_visibility="collapsed"
    )
    
    if uploaded_file is not None:
        try:
            loaded_data = json.loads(uploaded_file.read().decode('utf-8'))
            
            if st.button("불러오기 적용", use_container_width=True):
                for key in ['topic', 'target_persona', 'pain_points', 'one_line_concept', 
                           'outline', 'chapters', 'book_title', 'subtitle', 'market_analysis',
                           'topic_score', 'topic_verdict', 'score_details', 'generated_titles']:
                    if key in loaded_data:
                        st.session_state[key] = loaded_data[key]
                
                st.success("불러오기 완료!")
                st.rerun()
        except Exception as e:
            st.error(f"파일 오류: {e}")
    
    st.markdown("---")
    st.markdown("### API 설정")
    
    if 'api_key' not in st.session_state:
        saved_key = load_saved_api_key()
        st.session_state['api_key'] = saved_key
    
    api_key_input = st.text_input(
        "Gemini API 키",
        value=st.session_state['api_key'],
        type="password",
        placeholder="AIza...",
        help="Google AI Studio에서 발급받은 API 키를 입력하세요"
    )
    
    if api_key_input and api_key_input != st.session_state['api_key']:
        st.session_state['api_key'] = api_key_input
        if save_api_key(api_key_input):
            st.toast("✅ API 키가 저장되었습니다!", icon="💾")
    elif api_key_input:
        st.session_state['api_key'] = api_key_input
    
    with st.expander("API 키 발급 방법 (무료)"):
        st.markdown("""
        **2분이면 끝!**
        
        1. [Google AI Studio](https://aistudio.google.com/apikey) 접속
        2. Google 계정으로 로그인
        3. **"API 키 만들기"** 클릭
        4. 생성된 키 복사
        5. 위 입력창에 붙여넣기
        
        ✅ 완전 무료  
        ✅ 신용카드 불필요  
        ✅ 분당 15회 요청 가능
        """)
    
    if not st.session_state.get('api_key'):
        st.caption("⚠️ API 키를 입력하세요")
    else:
        col_status, col_del = st.columns([3, 1])
        with col_status:
            st.caption("✅ API 키 입력됨 (자동 저장)")
        with col_del:
            if st.button("🗑️", key="del_api_key", help="API 키 삭제"):
                st.session_state['api_key'] = ''
                save_api_key('')
                st.rerun()

# --- AI 함수 ---
def get_api_key():
    return st.session_state.get('api_key', '')

def get_auto_save_data():
    """자동 저장용 데이터 생성"""
    return {
        'topic': st.session_state.get('topic', ''),
        'target_persona': st.session_state.get('target_persona', ''),
        'pain_points': st.session_state.get('pain_points', ''),
        'one_line_concept': st.session_state.get('one_line_concept', ''),
        'outline': st.session_state.get('outline', []),
        'chapters': st.session_state.get('chapters', {}),
        'book_title': st.session_state.get('book_title', ''),
        'subtitle': st.session_state.get('subtitle', ''),
        'market_analysis': st.session_state.get('market_analysis', ''),
        'topic_score': st.session_state.get('topic_score'),
        'topic_verdict': st.session_state.get('topic_verdict'),
        'score_details': st.session_state.get('score_details'),
        'generated_titles': st.session_state.get('generated_titles'),
        'saved_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }

def sync_full_outline():
    """현재 session_state의 outline과 chapters를 기반으로 full_outline 재생성"""
    if not st.session_state.get('outline'):
        return
    
    new_full_outline = ""
    for ch in st.session_state['outline']:
        new_full_outline += f"## {ch}\n"
        if ch in st.session_state.get('chapters', {}):
            for st_name in st.session_state['chapters'][ch].get('subtopics', []):
                new_full_outline += f"- {st_name}\n"
        new_full_outline += "\n"
    
    st.session_state['full_outline'] = new_full_outline.strip()

def trigger_auto_save():
    """자동 저장 트리거"""
    sync_full_outline()
    st.session_state['auto_save_trigger'] = True

def ask_ai(system_role, prompt, temperature=0.7):
    api_key = get_api_key()
    if not api_key:
        return "⚠️ API 키를 먼저 입력해주세요."
    
    try:
        genai.configure(api_key=api_key)
        ai_model = genai.GenerativeModel('models/gemini-2.0-flash')
        generation_config = genai.types.GenerationConfig(temperature=temperature)
        full_prompt = f"""당신은 {system_role}입니다.

{prompt}

반드시 한국어로만 답변하세요. 영어, 러시아어, 아랍어 등 다른 언어는 절대 사용하지 마세요."""
        response = ai_model.generate_content(full_prompt, generation_config=generation_config)
        return response.text
    except Exception as e:
        return f"오류 발생: {str(e)}"

def analyze_topic_score(topic):
    prompt = f"""'{topic}' 주제의 전자책 적합도를 분석해주세요.

다음 5가지 항목을 각각 0~100점으로 채점하고, 종합 점수와 판정을 내려주세요.

채점 항목:
1. 시장성 (수요가 있는가?)
2. 수익성 (돈을 지불할 의향이 있는 주제인가?)
3. 차별화 가능성 (경쟁에서 이길 수 있는가?)
4. 작성 난이도 (전자책으로 만들기 쉬운가?)
5. 지속성 (오래 팔릴 수 있는가?)

반드시 아래 JSON 형식으로만 답변하세요. 다른 텍스트 없이 JSON만:
{{
    "market": {{"score": 85, "reason": "이유"}},
    "profit": {{"score": 80, "reason": "이유"}},
    "differentiation": {{"score": 75, "reason": "이유"}},
    "difficulty": {{"score": 90, "reason": "이유"}},
    "sustainability": {{"score": 70, "reason": "이유"}},
    "total_score": 80,
    "verdict": "적합" 또는 "보통" 또는 "부적합",
    "summary": "종합 의견 2~3문장"
}}"""
    return ask_ai("전자책 시장 분석가", prompt, temperature=0.3)

def generate_titles_advanced(topic, persona, pain_points):
    prompt = f"""당신은 자청(역행자), 엠제이 드마코(부의 추월차선), 김승호(돈의 속성)급 베스트셀러 작가입니다.

[분석 대상]
주제: {topic}
타겟: {persona}  
타겟의 속마음: {pain_points}

[베스트셀러 제목의 핵심 원칙]

1. 기존 상식을 정면으로 뒤집어라
2. 소외감과 긴급함을 동시에 자극하라
3. 구체적 숫자는 신뢰를 만든다
4. 짧을수록 강하다 - 7자 이내 메인 타이틀

[절대 금지]
- "비법", "노하우", "성공", "시작하세요", "방법", "전략", "가이드"
- "~하는 법", "~하기", "완벽한", "쉬운", "단계별"

형식 (JSON만 출력):
{{
    "titles": [
        {{
            "title": "7자 이내 임팩트 제목",
            "subtitle": "15자 이내 보조 설명",
            "concept": "이 제목의 핵심 컨셉",
            "why_works": "왜 사람들이 이 제목에 끌리는지"
        }}
    ]
}}"""
    return ask_ai("베스트셀러 작가", prompt, temperature=0.9)

def generate_concept(topic, persona, pain_points):
    prompt = f"""주제: {topic}
타겟: {persona}
타겟의 고민: {pain_points}

"이 책 안 읽으면 손해"라는 느낌을 주는 한 줄 컨셉 5개를 만들어주세요.

출력 형식:

1. [한 줄 컨셉]
   → 왜 끌리는가

2. [한 줄 컨셉]
   → 왜 끌리는가

3. [한 줄 컨셉]
   → 왜 끌리는가

4. [한 줄 컨셉]
   → 왜 끌리는가

5. [한 줄 컨셉]
   → 왜 끌리는가"""
    return ask_ai("카피라이터", prompt, temperature=0.9)

def generate_outline(topic, persona, pain_points):
    prompt = f"""주제: {topic}
타겟: {persona}
타겟의 고민: {pain_points}

위 주제로 정확히 4개 챕터 목차를 설계해주세요.
각 챕터당 3개 소제목입니다. (총 4챕터 × 3소제목 = 12개)

[목차 작성 핵심 규칙]

1. 챕터 제목: 7글자 이내
   - 짧고 강렬하게
   - "~하는 법", "~의 비밀" 같은 설명형 금지
   - 좋은 예: "바닥의 맛", "첫 100만원", "그날 이후", "역전의 시작"
   - 나쁜 예: "부자가 되는 마인드셋의 비밀", "성공하는 투자 전략"

2. 소제목: 10글자 이내
   - 궁금증 유발, 클릭하고 싶게
   - 뻔한 내용 암시 금지
   - 좋은 예: "47만원의 바닥", "새벽 4시의 선택", "그 한 마디", "월급의 배신"
   - 나쁜 예: "투자의 중요성", "첫 번째 실수와 교훈", "성공적인 마인드셋"

3. 절대 금지
   - "~하는 방법", "~의 중요성", "~을 위한 전략"
   - "첫 번째", "두 번째" 나열
   - 내용을 설명하는 제목

출력 형식 (정확히 이 형식만):

## 챕터1: [7자 이내 제목]
- [10자 이내 소제목]
- [10자 이내 소제목]
- [10자 이내 소제목]

## 챕터2: [7자 이내 제목]
- [10자 이내 소제목]
- [10자 이내 소제목]
- [10자 이내 소제목]

## 챕터3: [7자 이내 제목]
- [10자 이내 소제목]
- [10자 이내 소제목]
- [10자 이내 소제목]

## 챕터4: [7자 이내 제목]
- [10자 이내 소제목]
- [10자 이내 소제목]
- [10자 이내 소제목]"""
    return ask_ai("출판기획자", prompt, temperature=0.9)

def regenerate_chapter_outline(chapter_number, topic, persona, existing_chapters):
    prompt = f"""주제: {topic}

{chapter_number}번째 챕터를 새로 만들어주세요.

[규칙]
- 챕터 제목: 7글자 이내, 강렬하게
- 소제목: 10글자 이내, 궁금증 유발
- 설명형 금지 ("~하는 법", "~의 비밀")

출력 (정확히 이 형식만):
## 챕터{chapter_number}: [7자 이내]
- [10자 이내]
- [10자 이내]
- [10자 이내]"""
    return ask_ai("출판기획자", prompt, temperature=0.9)

def regenerate_single_subtopic(chapter_title, subtopic_index, topic, existing_subtopics):
    prompt = f"""주제: {topic}
챕터: {chapter_title}

{subtopic_index}번 소제목을 새로 만들어주세요.
10자 이내, 궁금증 유발하는 제목으로.
설명형 금지 ("~하는 법", "~의 중요성")

출력 (소제목만):"""
    result = ask_ai("출판기획자", prompt, temperature=0.9)
    result = result.strip().strip('[]').strip('-').strip()
    if '\n' in result:
        result = result.split('\n')[0].strip()
    return result

def generate_subtopics(chapter_title, topic, persona, num_subtopics=3):
    prompt = f"""주제: {topic}
챕터: {chapter_title}
타겟: {persona}

이 챕터의 소제목 {num_subtopics}개를 만들어주세요.

[규칙]
- 10글자 이내
- 궁금증 유발, 클릭하고 싶게
- 설명형 금지 ("~하는 법", "~의 중요성")

출력 형식 (숫자와 소제목만):
1. 소제목
2. 소제목
3. 소제목"""
    return ask_ai("베스트셀러 작가", prompt, temperature=0.95)

def generate_interview_questions(subtopic_title, chapter_title, topic):
    prompt = f"""'{topic}' 전자책의 '{chapter_title}' 챕터 중 '{subtopic_title}' 소제목 부분을 쓰기 위해 작가를 인터뷰합니다.

'{subtopic_title}'에 대한 작가의 진짜 경험과 통찰을 끌어내는 질문 3개를 만들어주세요.

형식:
Q1: [구체적이고 깊이 있는 질문]
Q2: [구체적이고 깊이 있는 질문]
Q3: [구체적이고 깊이 있는 질문]"""
    return ask_ai("베스트셀러 고스트라이터", prompt, temperature=0.7)


# ==========================================
# 🔥 핵심 수정: PS글쓰기 스타일 본문 생성 함수
# ==========================================
def generate_subtopic_content(subtopic_title, chapter_title, questions, answers, topic, persona):
    """프드프/자청 스타일 본문 생성 - 몰입감, 문단 구성, 스토리텔링"""
    
    qa_pairs = ""
    for i, (q, a) in enumerate(zip(questions, answers), 1):
        if a.strip():
            qa_pairs += f"\n질문{i}: {q}\n답변{i}: {a}\n"
    
    prompt = f"""당신은 '프드프', '자청'처럼 몰입감 높은 전자책을 쓰는 베스트셀러 작가입니다.

[집필 정보]
주제: {topic}
챕터: {chapter_title}
소제목: {subtopic_title}
타겟: {persona}

[작가 인터뷰 - 이 내용을 바탕으로]
{qa_pairs}

===================================
🔥 가장 중요: 문단 구성 규칙
===================================

[반드시 지켜야 할 양식]
- 한 문단 = 3~5문장을 붙여서 작성
- 문단과 문단 사이만 빈 줄 하나
- 한 문장마다 줄바꿈 절대 금지

[올바른 예시]
저는 그날 새벽 4시에 눈을 떴습니다. 통장 잔고는 47만원. 다음 달 월세를 내면 남는 건 없었습니다. 천장을 바라보며 생각했습니다. 이대로는 안 되겠다고.

그때 우연히 본 영상 하나가 제 인생을 바꿨습니다. 별거 아닌 내용이었습니다. 하지만 그 안에 제가 몰랐던 진실이 있었습니다. 돈을 버는 사람들은 뭔가 다르다는 걸 그제서야 알았습니다.

[잘못된 예시 - 이렇게 쓰지 마세요]
저는 그날 새벽 4시에 눈을 떴습니다.

통장 잔고는 47만원.

다음 달 월세를 내면 남는 건 없었습니다.

===================================
스토리텔링 규칙
===================================

1. 첫 문장은 강렬하게
   - 독자가 "뭐지?" 하고 멈추게
   - 결론이나 충격적 장면부터 시작
   - 예: "저는 그날 회사를 그만뒀습니다."

2. 구체적인 장면 묘사
   - 시간: "새벽 4시", "퇴근 후 지하철에서"
   - 장소: "원룸 책상 앞에서", "카페 구석자리에서"
   - 감정: "손이 떨렸습니다", "눈물이 났습니다"
   - 숫자: "47만원", "3개월", "월 300만원"

3. 독자 공감 포인트
   - "당신도 그런 적 있지 않나요?"라고 느끼게
   - 실패와 좌절의 순간을 솔직하게
   - 그때 느꼈던 감정을 생생하게

4. 스토리 흐름
   - 평범한 일상 → 문제 발생 → 고민과 시도 → 깨달음 → 변화
   - 교훈을 말하지 말고 이야기로 보여주세요

===================================
절대 금지 사항
===================================

- 한 문장씩 띄어쓰기 (가독성 파괴)
- "실수 1:", "해결책:" 같은 구조화
- "첫째,", "둘째," 나열
- "다음과 같습니다", "정리하면"
- "중요합니다", "필수입니다" 반복
- 번호 매기기(1. 2. 3.)
- 마크다운(**굵게**, *기울임*)
- 외국어(영어, 러시아어 등)
- 소제목/챕터 제목을 본문에서 반복
- 교훈을 직접 말하기 ("~해야 합니다")

===================================
[분량] 1500~2000자
[문체] 존댓말 통일 ("~입니다", "~습니다")
===================================

[미션]
'{subtopic_title}' 본문을 작성하세요.
- 소설처럼 몰입해서 읽히게
- 문단 단위로 구성 (3~5문장씩 묶어서)
- 독자가 "이건 내 얘기네" 하고 공감하게
- 끝까지 손에서 놓을 수 없게"""

    return ask_ai("베스트셀러 작가", prompt, temperature=0.75)


def refine_content(content, style="친근한"):
    style_guide = {
        "친근한": """친근한 스타일
- 존댓말로 통일 ("~입니다", "~합니다")
- 독자에게 직접 말하듯""",
        
        "전문적": """전문가 스타일
- 존댓말로 통일
- 데이터와 논리 강조""",
        
        "직설적": """직설 스타일
- 존댓말로 통일
- 핵심만 간결하게""",
        
        "스토리텔링": """스토리 스타일
- 존댓말로 통일
- 구체적 장면 묘사"""
    }
    
    prompt = f"""다음 글을 다듬어주세요.

[원본]
{content}

[가장 중요한 수정사항 - 문단 구성]
1. 한 문단 = 3~5문장 붙여서 작성
2. 한 문장마다 줄바꿈 절대 금지
3. 문단과 문단 사이만 빈 줄 하나

[올바른 예시]
저는 그날 새벽 4시에 눈을 떴습니다. 통장 잔고는 47만원. 다음 달 월세를 내면 남는 건 없었습니다. 천장을 바라보며 생각했습니다.

그때 우연히 본 영상 하나가 제 인생을 바꿨습니다. 별거 아닌 내용이었습니다. 하지만 그 안에 제가 몰랐던 진실이 있었습니다.

[추가 수정사항]
1. 존댓말로 통일 ("~입니다", "~습니다")
2. 반말("~하다", "~해라") → 존댓말로 변경
3. "실수 1:", "해결책:" 같은 구조화 표현 제거
4. 마크다운(**굵게**) 제거
5. 외국어 → 한국어로 번역

[목표 스타일]
{style_guide.get(style, style_guide["친근한"])}

[출력]
다듬어진 글만 출력. 설명 없이."""
    return ask_ai("에디터", prompt, temperature=0.7)

def check_quality(content):
    prompt = f"""다음 글이 베스트셀러 수준인지 평가해주세요.

[평가할 글]
{content[:4000]}

[평가 기준]
1. 첫 문장 (10점) - 독자의 뒤통수를 치는가?
2. 몰입도 (10점) - 끝까지 읽게 되는가?
3. 공감력 (10점) - "이건 내 얘기잖아"라고 느끼는가?
4. 구체성 (10점) - 구체적 장면/숫자가 있는가?
5. 문단 구성 (10점) - 3~5문장이 한 문단으로 잘 묶여 있는가? (한 문장씩 띄어쓰기는 감점)
6. AI 티 (10점) - AI 표현이 있는가?

[출력 형식]
📊 종합 점수: __/60점

📌 각 항목 평가와 개선안

✍️ 수정하면 좋을 문장 TOP 3

🎯 총평"""
    return ask_ai("베스트셀러 편집자", prompt, temperature=0.6)

def generate_marketing_copy(title, subtitle, topic, persona):
    prompt = f"""[상품 정보]
제목: {title}
부제: {subtitle}
주제: {topic}
타겟: {persona}

이 전자책을 폭발적으로 팔기 위한 킬러 카피를 만들어주세요.

1. 크몽 상품 제목 (40자 이내)
2. 상세페이지 헤드라인 3개
3. 구매 유도 문구 (CTA) 3개
4. 인스타그램 홍보 문구
5. 블로그 포스팅 제목 3개"""
    return ask_ai("크몽 탑셀러 마케터", prompt, temperature=0.85)


# ==========================================
# 글자 수 계산 헬퍼 함수
# ==========================================
def calculate_char_count(text):
    """순수 본문만으로 글자 수 계산 (공백, 줄바꿈 제외)"""
    if not text:
        return 0
    return len(text.replace('\n', '').replace(' ', ''))

def clean_content_for_display(content, subtopic_title=None, chapter_title=None):
    """본문에서 마크다운 기호, HTML 태그, 중복 제목 제거"""
    if not content:
        return ""
    
    # HTML 태그 제거
    content = re.sub(r'<[^>]+>', '', content)
    content = content.replace('&amp;', '&')
    content = content.replace('&lt;', '<')
    content = content.replace('&gt;', '>')
    content = content.replace('&quot;', '"')
    content = content.replace('&#39;', "'")
    content = content.replace('&nbsp;', ' ')
    
    lines = content.split('\n')
    cleaned_lines = []
    
    for idx, line in enumerate(lines):
        stripped = line.strip()
        
        if not stripped:
            if idx > 3 or len(cleaned_lines) > 0:
                cleaned_lines.append(line)
            continue
        
        # 마크다운 헤더 제거
        if stripped.startswith('#'):
            continue
        
        # "챕터 N:" 형식 제거
        if stripped.startswith('챕터') and ':' in stripped[:15]:
            continue
        
        # "소제목:" 형식 제거
        if stripped.startswith('소제목') and ':' in stripped[:10]:
            continue
        
        cleaned_lines.append(line)
    
    result = '\n'.join(cleaned_lines).strip()
    return result

def get_all_content_text():
    """모든 챕터의 순수 본문 텍스트만 수집"""
    pure_content = ""
    for ch in st.session_state.get('outline', []):
        if ch in st.session_state.get('chapters', {}):
            ch_data = st.session_state['chapters'][ch]
            if 'subtopic_data' in ch_data:
                subtopic_list = ch_data.get('subtopics', [])
                if not subtopic_list and ch in ch_data['subtopic_data']:
                    subtopic_list = [ch]
                for st_name in subtopic_list:
                    st_data = ch_data['subtopic_data'].get(st_name, {})
                    if st_data.get('content'):
                        pure_content += st_data['content']
    return pure_content


# --- 메인 UI ---
st.markdown("""
<div class="hero-section">
    <div class="hero-label">CASHMAKER</div>
    <div class="hero-title">전자책 작성 프로그램</div>
    <div class="hero-subtitle">쉽고, 빠른 전자책 수익화</div>
</div>
""", unsafe_allow_html=True)

# 메인 탭
tabs = st.tabs([
    "① 주제 선정", 
    "② 타겟 & 컨셉", 
    "③ 목차 설계", 
    "④ 본문 작성", 
    "⑤ 문체 다듬기",
    "⑥ 최종 출력"
])

# === TAB 1: 주제 선정 ===
with tabs[0]:
    st.markdown("## 주제 선정 & 적합도 분석")
    
    st.markdown("""
    <div class="quick-action-box">
        <p>💡 <strong>이미 주제가 있다면?</strong> 아래에 입력 후 바로 다음 탭으로 이동하세요!</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown('<p class="section-label">Step 01</p>', unsafe_allow_html=True)
        st.markdown("### 주제 입력")
        
        topic_input = st.text_input(
            "어떤 주제로 전자책을 쓰고 싶으세요?",
            value=st.session_state['topic'],
            placeholder="예: 크몽으로 월 500만원 벌기"
        )
        
        if topic_input != st.session_state['topic']:
            st.session_state['topic'] = topic_input
            st.session_state['topic_score'] = None
            st.session_state['score_details'] = None
        
        st.markdown("""
        <div class="info-card">
            <div class="info-card-title">좋은 주제의 조건</div>
            <p>• 내가 직접 경험하고 성과를 낸 것</p>
            <p>• 사람들이 돈 주고 배우고 싶어하는 것</p>
            <p>• 구체적인 결과를 약속할 수 있는 것</p>
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("📊 적합도 분석하기 (선택)", key="analyze_btn"):
            if not topic_input:
                st.error("주제를 입력해주세요.")
            else:
                with st.spinner("분석 중..."):
                    result = analyze_topic_score(topic_input)
                    try:
                        json_match = re.search(r'\{[\s\S]*\}', result)
                        if json_match:
                            score_data = json.loads(json_match.group())
                            st.session_state['topic_score'] = score_data.get('total_score', 0)
                            st.session_state['topic_verdict'] = score_data.get('verdict', '분석 실패')
                            st.session_state['score_details'] = score_data
                    except:
                        st.error("분석 결과 파싱 오류. 다시 시도해주세요.")
    
    with col2:
        st.markdown('<p class="section-label">Step 02</p>', unsafe_allow_html=True)
        st.markdown("### 분석 결과")
        
        if st.session_state['topic_score'] is not None:
            score = st.session_state['topic_score']
            verdict = st.session_state['topic_verdict']
            details = st.session_state['score_details']
            
            verdict_class = "status-excellent" if verdict == "적합" else ("status-good" if verdict == "보통" else "status-warning")
            
            st.markdown(f"""
            <div class="score-card">
                <div class="score-number">{score}</div>
                <div class="score-label">종합 점수</div>
                <span class="status-badge {verdict_class}">{verdict}</span>
            </div>
            """, unsafe_allow_html=True)
            
            if details:
                st.markdown("#### 세부 점수")
                
                items = [
                    ("시장성", details.get('market', {}).get('score', 0), details.get('market', {}).get('reason', '')),
                    ("수익성", details.get('profit', {}).get('score', 0), details.get('profit', {}).get('reason', '')),
                    ("차별화", details.get('differentiation', {}).get('score', 0), details.get('differentiation', {}).get('reason', '')),
                    ("작성 난이도", details.get('difficulty', {}).get('score', 0), details.get('difficulty', {}).get('reason', '')),
                    ("지속성", details.get('sustainability', {}).get('score', 0), details.get('sustainability', {}).get('reason', '')),
                ]
                
                for name, score_val, reason in items:
                    st.markdown(f"""
                    <div class="score-item">
                        <span class="score-item-label">{name}</span>
                        <span class="score-item-value">{score_val}</span>
                    </div>
                    <p class="score-item-reason">{reason}</p>
                    """, unsafe_allow_html=True)
                
                st.markdown(f"""
                <div class="summary-box">
                    <p><strong>종합 의견</strong><br>{details.get('summary', '')}</p>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="empty-state">
                <p>분석은 선택사항입니다.</p>
                <p>주제만 입력해도 다음 단계로 진행 가능!</p>
            </div>
            """, unsafe_allow_html=True)

# === TAB 2: 타겟 & 컨셉 ===
with tabs[1]:
    st.markdown("## 타겟 설정 & 제목 생성")
    
    if not st.session_state['topic']:
        st.info("💡 주제를 먼저 입력하면 더 정확한 결과를 얻을 수 있어요.")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown('<p class="section-label">Step 01</p>', unsafe_allow_html=True)
        st.markdown("### 타겟 정의")
        
        if not st.session_state['topic']:
            topic_here = st.text_input(
                "주제 (여기서 입력 가능)",
                value=st.session_state['topic'],
                placeholder="예: 크몽으로 월 500만원 벌기",
                key="topic_tab2"
            )
            if topic_here:
                st.session_state['topic'] = topic_here
        
        persona = st.text_area(
            "누가 이 책을 읽나요?",
            value=st.session_state['target_persona'],
            placeholder="예: 30대 직장인, 퇴근 후 부업으로 월 100만원 추가 수입을 원하는 사람",
            height=100
        )
        st.session_state['target_persona'] = persona
        
        pain_points = st.text_area(
            "타겟의 가장 큰 고민은?",
            value=st.session_state['pain_points'],
            placeholder="예: 시간이 없다, 뭘 해야 할지 모르겠다, 시작이 두렵다",
            height=100
        )
        st.session_state['pain_points'] = pain_points
        
        st.markdown("---")
        
        st.markdown('<p class="section-label">Step 02</p>', unsafe_allow_html=True)
        st.markdown("### 한 줄 컨셉")
        
        if st.button("컨셉 생성하기", key="concept_btn"):
            if not st.session_state['topic'] or not persona:
                st.error("주제와 타겟을 먼저 입력해주세요.")
            else:
                with st.spinner("생성 중..."):
                    concept = generate_concept(
                        st.session_state['topic'],
                        persona,
                        pain_points
                    )
                    st.session_state['one_line_concept'] = concept
        
        if st.session_state['one_line_concept']:
            st.markdown(f"""
            <div class="info-card">
                {st.session_state['one_line_concept'].replace(chr(10), '<br>')}
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        st.markdown('<p class="section-label">Step 03</p>', unsafe_allow_html=True)
        st.markdown("### 제목 생성")
        
        if st.button("제목 생성하기", key="title_btn"):
            if not st.session_state['topic']:
                st.error("주제를 먼저 입력해주세요.")
            else:
                with st.spinner("생성 중..."):
                    titles_result = generate_titles_advanced(
                        st.session_state['topic'],
                        st.session_state['target_persona'],
                        st.session_state['pain_points']
                    )
                    try:
                        json_match = re.search(r'\{[\s\S]*\}', titles_result)
                        if json_match:
                            st.session_state['generated_titles'] = json.loads(json_match.group())
                    except:
                        st.session_state['generated_titles'] = None
                        st.markdown(titles_result)
        
        if st.session_state.get('generated_titles'):
            titles_data = st.session_state['generated_titles']
            if 'titles' in titles_data:
                for i, t in enumerate(titles_data['titles'], 1):
                    st.markdown(f"""
                    <div class="title-card">
                        <div class="card-number">TITLE 0{i}</div>
                        <div class="main-title">{t.get('title', '')}</div>
                        <div class="sub-title">{t.get('subtitle', '')}</div>
                        <div class="reason">{t.get('why_works', '')}</div>
                    </div>
                    """, unsafe_allow_html=True)
        
        st.markdown("---")
        st.markdown('<p class="section-label">Step 04</p>', unsafe_allow_html=True)
        st.markdown("### 최종 선택")
        st.session_state['book_title'] = st.text_input("제목", value=st.session_state['book_title'], placeholder="최종 제목")
        st.session_state['subtitle'] = st.text_input("부제", value=st.session_state['subtitle'], placeholder="부제")

# === TAB 3: 목차 설계 ===
with tabs[2]:
    st.markdown("## 목차 설계")
    
    st.markdown("### 🎯 작업 방식 선택")
    outline_mode = st.radio(
        "목차를 어떻게 만드시겠어요?",
        ["🤖 자동으로 목차 생성", "✍️ 내가 직접 입력"],
        horizontal=True,
        key="outline_mode_radio"
    )
    
    st.markdown("---")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        if outline_mode == "🤖 자동으로 목차 생성":
            st.markdown('<p class="section-label">자동 목차 생성</p>', unsafe_allow_html=True)
            st.markdown("### 목차를 자동으로 설계합니다")
            
            if not st.session_state['topic']:
                st.warning("💡 주제를 먼저 입력해주세요")
                topic_here = st.text_input(
                    "주제",
                    value=st.session_state['topic'],
                    placeholder="예: 크몽으로 월 500만원 벌기",
                    key="topic_tab3"
                )
                if topic_here:
                    st.session_state['topic'] = topic_here
            
            if st.button("🚀 목차 생성하기", key="outline_btn"):
                if not st.session_state['topic']:
                    st.error("주제를 먼저 입력해주세요.")
                else:
                    with st.spinner("설계 중..."):
                        outline_text = generate_outline(
                            st.session_state['topic'],
                            st.session_state['target_persona'],
                            st.session_state['pain_points']
                        )
                        
                        lines = outline_text.split('\n')
                        chapters = []
                        current_chapter = None
                        chapter_subtopics = {}
                        
                        for line in lines:
                            line = line.strip()
                            if not line or line == '...':
                                continue
                            
                            def is_chapter_line_ai(text):
                                text_clean = text.lstrip('#').strip()
                                text_lower = text_clean.lower()
                                if any(text_lower.startswith(kw) for kw in ['챕터', 'chapter', '에필로그', '프롤로그', '서문', '부록']):
                                    return True
                                if len(text_clean) > 1 and text_clean[0].isdigit():
                                    rest = text_clean[1:].lstrip('0123456789')
                                    if rest and (rest[0] in '부장.:'):
                                        return True
                                return False
                            
                            def is_subtopic_line_ai(text):
                                if text[0] in '-·•':
                                    return True
                                if len(text) > 1 and text[0].isdigit():
                                    for i, char in enumerate(text):
                                        if char == ')':
                                            return True
                                        if not char.isdigit():
                                            break
                                return False
                            
                            if is_chapter_line_ai(line):
                                chapter_name = line.lstrip('#').strip()
                                current_chapter = chapter_name
                                chapters.append(current_chapter)
                                chapter_subtopics[current_chapter] = []
                            elif current_chapter:
                                if is_subtopic_line_ai(line):
                                    subtopic = line.lstrip('-·• ')
                                    subtopic = re.sub(r'^\d+\)\s*', '', subtopic)
                                    if subtopic:
                                        chapter_subtopics[current_chapter].append(subtopic)
                        
                        st.session_state['outline'] = chapters
                        st.session_state['full_outline'] = outline_text
                        
                        for ch in chapters:
                            subtopics = chapter_subtopics.get(ch, [])
                            st.session_state['chapters'][ch] = {
                                'subtopics': subtopics,
                                'subtopic_data': {st: {'questions': [], 'answers': [], 'content': ''} for st in subtopics}
                            }
                        
                        total_subtopics = sum(len(chapter_subtopics.get(ch, [])) for ch in chapters)
                        st.success(f"✅ {len(chapters)}개 챕터, {total_subtopics}개 소제목 생성됨!")
                        st.rerun()
            
            if 'full_outline' in st.session_state and st.session_state['full_outline']:
                st.markdown("**📋 현재 목차**")
                st.code(st.session_state['full_outline'], language=None)
        
        else:
            st.markdown('<p class="section-label">직접 입력</p>', unsafe_allow_html=True)
            st.markdown("### 목차를 직접 입력하세요")
            
            existing_outline = ""
            if st.session_state['outline']:
                for ch in st.session_state['outline']:
                    existing_outline += f"{ch}\n"
                    if ch in st.session_state['chapters']:
                        for i, st_name in enumerate(st.session_state['chapters'][ch].get('subtopics', []), 1):
                            existing_outline += f"{i}) {st_name}\n"
            
            manual_outline = st.text_area(
                "목차 입력 (챕터와 소제목)",
                value=existing_outline,
                height=350,
                placeholder="1부. 첫 번째 챕터 제목\n1) 소제목 1\n2) 소제목 2\n...",
                key="manual_outline_input"
            )
            
            if st.button("✅ 목차 저장하기", key="save_manual_outline"):
                if manual_outline.strip():
                    lines = manual_outline.strip().split('\n')
                    chapters = []
                    current_chapter = None
                    chapter_subtopics = {}
                    
                    for line in lines:
                        line = line.strip()
                        if not line or line == '...':
                            continue
                        
                        is_chapter = False
                        text_lower = line.lower()
                        
                        if any(text_lower.startswith(kw) for kw in ['챕터', 'chapter', '에필로그', '프롤로그', '서문', '부록']):
                            is_chapter = True
                        elif len(line) > 2 and line[0].isdigit():
                            num_end = 1
                            while num_end < len(line) and line[num_end].isdigit():
                                num_end += 1
                            rest = line[num_end:]
                            
                            if rest:
                                if rest[0] in '부장':
                                    is_chapter = True
                                elif rest[0] == '.' and len(rest) > 1 and rest[1] != ')':
                                    is_chapter = True
                                elif rest[0] == ':':
                                    is_chapter = True
                        
                        is_subtopic = False
                        if not is_chapter and current_chapter:
                            if line[0] in '-·•':
                                is_subtopic = True
                            elif line[0].isdigit():
                                num_end = 1
                                while num_end < len(line) and line[num_end].isdigit():
                                    num_end += 1
                                if num_end < len(line) and line[num_end] == ')':
                                    is_subtopic = True
                        
                        if is_chapter:
                            current_chapter = line
                            chapters.append(current_chapter)
                            chapter_subtopics[current_chapter] = []
                        elif current_chapter:
                            subtopic = line
                            if line[0] in '-·•':
                                subtopic = line.lstrip('-·• ').strip()
                            elif is_subtopic:
                                subtopic = re.sub(r'^\d+\)\s*', '', line).strip()
                            
                            if subtopic and len(subtopic) > 2:
                                chapter_subtopics[current_chapter].append(subtopic)
                    
                    st.session_state['outline'] = chapters
                    st.session_state['full_outline'] = manual_outline
                    
                    for ch in chapters:
                        subtopics = chapter_subtopics.get(ch, [])
                        st.session_state['chapters'][ch] = {
                            'subtopics': subtopics,
                            'subtopic_data': {st_name: {'questions': [], 'answers': [], 'content': ''} for st_name in subtopics}
                        }
                    
                    trigger_auto_save()
                    
                    total_subtopics = sum(len(chapter_subtopics.get(ch, [])) for ch in chapters)
                    st.success(f"✅ {len(chapters)}개 챕터, {total_subtopics}개 소제목 저장됨!")
                    st.rerun()
                else:
                    st.error("목차를 입력해주세요.")
    
    with col2:
        st.markdown('<p class="section-label">목차 관리</p>', unsafe_allow_html=True)
        st.markdown("### 📋 현재 목차")
        
        if st.session_state['outline']:
            for i, chapter in enumerate(st.session_state['outline']):
                subtopic_count = 0
                if chapter in st.session_state['chapters']:
                    subtopic_count = len(st.session_state['chapters'][chapter].get('subtopics', []))
                
                with st.expander(f"**{chapter}** ({subtopic_count}개 소제목)", expanded=False):
                    col_edit, col_actions = st.columns([3, 2])
                    with col_edit:
                        new_title = st.text_input(
                            "챕터 제목",
                            value=chapter,
                            key=f"edit_chapter_{i}",
                            label_visibility="collapsed"
                        )
                    with col_actions:
                        col_regen, col_del = st.columns(2)
                        with col_regen:
                            if st.button("🔄", key=f"regen_chapter_{i}", help="이 챕터만 새로 생성"):
                                with st.spinner("챕터 재생성 중..."):
                                    new_chapter_text = regenerate_chapter_outline(
                                        i + 1,
                                        st.session_state['topic'],
                                        st.session_state['target_persona'],
                                        st.session_state['outline']
                                    )
                                    lines = new_chapter_text.split('\n')
                                    new_chapter_title = None
                                    new_subtopics = []
                                    for line in lines:
                                        line = line.strip()
                                        if line.startswith('##'):
                                            new_chapter_title = line.lstrip('#').strip()
                                        elif line.startswith('-'):
                                            st_name = line.lstrip('- ').strip()
                                            if st_name:
                                                new_subtopics.append(st_name)
                                    
                                    if new_chapter_title:
                                        old_chapter = st.session_state['outline'][i]
                                        st.session_state['outline'][i] = new_chapter_title
                                        if old_chapter in st.session_state['chapters']:
                                            del st.session_state['chapters'][old_chapter]
                                        st.session_state['chapters'][new_chapter_title] = {
                                            'subtopics': new_subtopics,
                                            'subtopic_data': {st: {'questions': [], 'answers': [], 'content': ''} for st in new_subtopics}
                                        }
                                        trigger_auto_save()
                                        st.rerun()
                        with col_del:
                            if st.button("🗑️", key=f"del_chapter_{i}", help="삭제"):
                                old_chapter = st.session_state['outline'].pop(i)
                                if old_chapter in st.session_state['chapters']:
                                    del st.session_state['chapters'][old_chapter]
                                trigger_auto_save()
                                st.rerun()
                    
                    if new_title != chapter and new_title.strip():
                        if st.button("💾 제목 저장", key=f"save_chapter_title_{i}"):
                            st.session_state['outline'][i] = new_title
                            if chapter in st.session_state['chapters']:
                                st.session_state['chapters'][new_title] = st.session_state['chapters'].pop(chapter)
                            trigger_auto_save()
                            st.rerun()
                    
                    st.markdown("---")
                    st.markdown("**📝 소제목 관리**")
                    
                    if chapter in st.session_state['chapters']:
                        subtopics = st.session_state['chapters'][chapter].get('subtopics', [])
                        
                        for j, st_name in enumerate(subtopics):
                            col_st, col_st_actions = st.columns([3, 2])
                            with col_st:
                                new_st = st.text_input(
                                    f"소제목 {j+1}",
                                    value=st_name,
                                    key=f"edit_st_{i}_{j}",
                                    label_visibility="collapsed"
                                )
                            with col_st_actions:
                                col_st_save, col_st_regen, col_st_del = st.columns(3)
                                with col_st_save:
                                    if new_st != st_name and new_st.strip():
                                        if st.button("💾", key=f"save_st_{i}_{j}", help="저장"):
                                            st.session_state['chapters'][chapter]['subtopics'][j] = new_st
                                            if st_name in st.session_state['chapters'][chapter]['subtopic_data']:
                                                st.session_state['chapters'][chapter]['subtopic_data'][new_st] = st.session_state['chapters'][chapter]['subtopic_data'].pop(st_name)
                                            trigger_auto_save()
                                            st.rerun()
                                with col_st_regen:
                                    if st.button("🔄", key=f"regen_st_{i}_{j}", help="재생성"):
                                        with st.spinner("소제목 재생성 중..."):
                                            new_st_title = regenerate_single_subtopic(
                                                chapter,
                                                j + 1,
                                                st.session_state['topic'],
                                                subtopics
                                            )
                                            if new_st_title:
                                                old_st = st.session_state['chapters'][chapter]['subtopics'][j]
                                                st.session_state['chapters'][chapter]['subtopics'][j] = new_st_title
                                                if old_st in st.session_state['chapters'][chapter]['subtopic_data']:
                                                    st.session_state['chapters'][chapter]['subtopic_data'][new_st_title] = st.session_state['chapters'][chapter]['subtopic_data'].pop(old_st)
                                                else:
                                                    st.session_state['chapters'][chapter]['subtopic_data'][new_st_title] = {'questions': [], 'answers': [], 'content': ''}
                                                trigger_auto_save()
                                                st.rerun()
                                with col_st_del:
                                    if st.button("🗑️", key=f"del_st_{i}_{j}", help="삭제"):
                                        removed_st = st.session_state['chapters'][chapter]['subtopics'].pop(j)
                                        if removed_st in st.session_state['chapters'][chapter]['subtopic_data']:
                                            del st.session_state['chapters'][chapter]['subtopic_data'][removed_st]
                                        trigger_auto_save()
                                        st.rerun()
                        
                        st.markdown("---")
                        col_add_st, col_add_btn = st.columns([3, 1])
                        with col_add_st:
                            new_st_input = st.text_input(
                                "새 소제목",
                                placeholder="새 소제목을 입력하세요",
                                key=f"new_st_input_{i}",
                                label_visibility="collapsed"
                            )
                        with col_add_btn:
                            if st.button("➕", key=f"add_st_{i}"):
                                if new_st_input.strip():
                                    st.session_state['chapters'][chapter]['subtopics'].append(new_st_input.strip())
                                    st.session_state['chapters'][chapter]['subtopic_data'][new_st_input.strip()] = {
                                        'questions': [], 'answers': [], 'content': ''
                                    }
                                    trigger_auto_save()
                                    st.rerun()
                    else:
                        st.info("소제목이 없습니다. 아래에서 추가하세요.")
            
            st.markdown("---")
            
            if st.button("➕ 새 챕터 추가", key="add_chapter"):
                new_ch_name = f"챕터{len(st.session_state['outline'])+1}: 새 챕터"
                st.session_state['outline'].append(new_ch_name)
                st.session_state['chapters'][new_ch_name] = {
                    'subtopics': [],
                    'subtopic_data': {}
                }
                trigger_auto_save()
                st.rerun()
            
        else:
            st.markdown("""
            <div class="empty-state">
                <p>왼쪽에서 목차를 생성하거나 직접 입력하세요</p>
            </div>
            """, unsafe_allow_html=True)

# === TAB 4: 본문 작성 ===
with tabs[3]:
    st.markdown("## 본문 작성")
    
    if not st.session_state['outline']:
        st.warning("⚠️ 먼저 '③ 목차 설계' 탭에서 목차를 작성해주세요.")
        st.stop()
    
    chapter_list = []
    for item in st.session_state['outline']:
        item_stripped = item.strip()
        if not item_stripped.startswith('-') and not item_stripped.startswith('·') and not item_stripped.startswith('•'):
            chapter_list.append(item)
    
    if not chapter_list:
        st.warning("⚠️ 챕터가 없습니다. 목차를 다시 확인해주세요.")
        st.stop()
    
    selected_chapter = st.selectbox(
        "📚 챕터 선택",
        chapter_list,
        key="chapter_select_main"
    )
    
    if selected_chapter not in st.session_state['chapters']:
        st.session_state['chapters'][selected_chapter] = {
            'subtopics': [],
            'subtopic_data': {}
        }
    
    chapter_data = st.session_state['chapters'][selected_chapter]
    
    if 'subtopics' not in chapter_data:
        chapter_data['subtopics'] = []
    if 'subtopic_data' not in chapter_data:
        chapter_data['subtopic_data'] = {}
    
    for st_name in chapter_data['subtopics']:
        if st_name not in chapter_data['subtopic_data']:
            chapter_data['subtopic_data'][st_name] = {'questions': [], 'answers': [], 'content': ''}
    
    st.markdown("---")
    
    with st.expander(f"📋 '{selected_chapter}' 소제목 전체 보기 ({len(chapter_data.get('subtopics', []))}개)", expanded=False):
        if chapter_data.get('subtopics'):
            for j, st_name in enumerate(chapter_data['subtopics']):
                has_content = bool(chapter_data['subtopic_data'].get(st_name, {}).get('content', '').strip())
                status_icon = "✅" if has_content else "⬜"
                
                col_st_view, col_st_edit, col_st_regen = st.columns([4, 1, 1])
                with col_st_view:
                    new_st_name = st.text_input(
                        f"{status_icon} {j+1}",
                        value=st_name,
                        key=f"view_st_tab4_{j}",
                        label_visibility="collapsed"
                    )
                with col_st_edit:
                    if new_st_name != st_name and new_st_name.strip():
                        if st.button("💾", key=f"save_st_tab4_{j}", help="저장"):
                            chapter_data['subtopics'][j] = new_st_name
                            if st_name in chapter_data['subtopic_data']:
                                chapter_data['subtopic_data'][new_st_name] = chapter_data['subtopic_data'].pop(st_name)
                            trigger_auto_save()
                            st.rerun()
                with col_st_regen:
                    if st.button("🔄", key=f"regen_st_tab4_{j}", help="이 소제목만 재생성"):
                        with st.spinner("재생성 중..."):
                            new_title = regenerate_single_subtopic(
                                selected_chapter,
                                j + 1,
                                st.session_state['topic'],
                                chapter_data['subtopics']
                            )
                            if new_title:
                                old_st = chapter_data['subtopics'][j]
                                chapter_data['subtopics'][j] = new_title
                                if old_st in chapter_data['subtopic_data']:
                                    chapter_data['subtopic_data'][new_title] = chapter_data['subtopic_data'].pop(old_st)
                                else:
                                    chapter_data['subtopic_data'][new_title] = {'questions': [], 'answers': [], 'content': ''}
                                trigger_auto_save()
                                st.rerun()
            
            st.markdown("---")
            col_new_st, col_new_btn = st.columns([4, 1])
            with col_new_st:
                new_st_input = st.text_input(
                    "새 소제목 추가",
                    placeholder="새 소제목을 입력하세요",
                    key="new_st_input_tab4",
                    label_visibility="collapsed"
                )
            with col_new_btn:
                if st.button("➕", key="add_st_tab4", help="추가"):
                    if new_st_input.strip():
                        chapter_data['subtopics'].append(new_st_input.strip())
                        chapter_data['subtopic_data'][new_st_input.strip()] = {
                            'questions': [], 'answers': [], 'content': ''
                        }
                        trigger_auto_save()
                        st.rerun()
        else:
            st.info("소제목이 없습니다.")
    
    st.markdown("---")
    
    if chapter_data['subtopics']:
        st.markdown("### ✍️ 소제목 선택 → 본문 작성")
        
        selected_subtopic = st.selectbox(
            "작성할 소제목",
            chapter_data['subtopics'],
            key="subtopic_select_main",
            format_func=lambda x: f"{'✅' if chapter_data['subtopic_data'].get(x, {}).get('content') else '⬜'} {x}"
        )
        
        completed = sum(1 for s in chapter_data['subtopics'] if chapter_data['subtopic_data'].get(s, {}).get('content'))
        total = len(chapter_data['subtopics'])
        st.progress(completed / total if total > 0 else 0)
        st.caption(f"진행: {completed}/{total} 완료")
        
        st.markdown("---")
        
        if selected_subtopic:
            if selected_subtopic not in chapter_data['subtopic_data']:
                chapter_data['subtopic_data'][selected_subtopic] = {'questions': [], 'answers': [], 'content': ''}
            
            subtopic_data = chapter_data['subtopic_data'][selected_subtopic]
            
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.markdown('<p class="section-label">Step 01</p>', unsafe_allow_html=True)
                st.markdown(f"### 🎤 인터뷰: {selected_subtopic}")
                
                if st.button("🎤 질문 생성하기", key="gen_questions_main"):
                    with st.spinner("질문 생성 중..."):
                        questions_text = generate_interview_questions(
                            selected_subtopic, 
                            selected_chapter, 
                            st.session_state['topic']
                        )
                        questions = re.findall(r'Q\d+:\s*(.+)', questions_text)
                        if not questions:
                            questions = [q.strip() for q in questions_text.split('\n') if q.strip() and '?' in q][:3]
                        subtopic_data['questions'] = questions
                        subtopic_data['answers'] = [''] * len(questions)
                        st.rerun()
                
                if subtopic_data['questions']:
                    for i, q in enumerate(subtopic_data['questions']):
                        st.markdown(f"**Q{i+1}.** {q}")
                        if i >= len(subtopic_data['answers']):
                            subtopic_data['answers'].append('')
                        subtopic_data['answers'][i] = st.text_area(
                            f"A{i+1}",
                            value=subtopic_data['answers'][i],
                            key=f"answer_main_{selected_chapter}_{selected_subtopic}_{i}",
                            height=80,
                            label_visibility="collapsed"
                        )
                else:
                    st.info("👆 '질문 생성하기' 버튼을 눌러 인터뷰를 시작하세요.")
            
            with col2:
                st.markdown('<p class="section-label">Step 02</p>', unsafe_allow_html=True)
                st.markdown(f"### 📝 본문: {selected_subtopic}")
                
                has_answers = subtopic_data.get('questions') and any(a.strip() for a in subtopic_data.get('answers', []))
                
                content_widget_key = f"content_main_{selected_chapter}_{selected_subtopic}"
                
                if has_answers:
                    if st.button("✨ 본문 생성하기", key="gen_content_main"):
                        with st.spinner("집필 중... (30초~1분)"):
                            content = generate_subtopic_content(
                                selected_subtopic,
                                selected_chapter,
                                subtopic_data['questions'],
                                subtopic_data['answers'],
                                st.session_state['topic'],
                                st.session_state['target_persona']
                            )
                            st.session_state['chapters'][selected_chapter]['subtopic_data'][selected_subtopic]['content'] = content
                            st.session_state[content_widget_key] = content
                            trigger_auto_save()
                            st.rerun()
                else:
                    st.info("👈 먼저 인터뷰 질문에 답변해주세요.")
                
                stored_content = st.session_state['chapters'][selected_chapter]['subtopic_data'][selected_subtopic].get('content', '')
                if content_widget_key not in st.session_state:
                    st.session_state[content_widget_key] = stored_content
                
                edited_content = st.text_area(
                    "본문 내용",
                    height=400,
                    key=content_widget_key,
                    label_visibility="collapsed"
                )
                
                if content_widget_key in st.session_state:
                    st.session_state['chapters'][selected_chapter]['subtopic_data'][selected_subtopic]['content'] = st.session_state[content_widget_key]
                
                final_content = st.session_state['chapters'][selected_chapter]['subtopic_data'][selected_subtopic].get('content', '')
                if final_content:
                    char_count = calculate_char_count(final_content)
                    st.caption(f"📊 {char_count:,}자")
                    st.success(f"✅ '{selected_subtopic}' 본문 작성 완료!")
        
        with st.expander("⚙️ 소제목 편집/추가", expanded=False):
            st.markdown("#### 소제목 관리")
            
            col_gen, col_add = st.columns(2)
            
            with col_gen:
                num_subtopics = st.number_input(
                    "생성할 개수",
                    min_value=1,
                    max_value=10,
                    value=3,
                    key="num_subtopics_gen_exp"
                )
                if st.button("✨ 소제목 자동 생성", key="gen_subtopics_exp"):
                    with st.spinner("생성 중..."):
                        subtopics_text = generate_subtopics(
                            selected_chapter,
                            st.session_state['topic'],
                            st.session_state['target_persona'],
                            num_subtopics
                        )
                        new_subtopics = []
                        for line in subtopics_text.split('\n'):
                            line = line.strip()
                            if line and (line[0].isdigit() or line.startswith('-')):
                                cleaned = re.sub(r'^[\d\.\-\s]+', '', line).strip()
                                if cleaned:
                                    new_subtopics.append(cleaned)
                        
                        if new_subtopics:
                            chapter_data['subtopics'] = new_subtopics[:num_subtopics]
                            for st_name in new_subtopics[:num_subtopics]:
                                if st_name not in chapter_data['subtopic_data']:
                                    chapter_data['subtopic_data'][st_name] = {'questions': [], 'answers': [], 'content': ''}
                            st.success(f"✅ {len(new_subtopics[:num_subtopics])}개 생성됨!")
                            st.rerun()
            
            with col_add:
                new_name = st.text_input("새 소제목", placeholder="직접 입력", key="new_subtopic_exp")
                if st.button("➕ 추가", key="add_subtopic_exp"):
                    if new_name.strip() and new_name not in chapter_data['subtopics']:
                        chapter_data['subtopics'].append(new_name)
                        chapter_data['subtopic_data'][new_name] = {'questions': [], 'answers': [], 'content': ''}
                        st.rerun()
            
            st.markdown("**현재 소제목:**")
            for i, st_name in enumerate(chapter_data['subtopics']):
                col_n, col_del = st.columns([5, 1])
                with col_n:
                    st.write(f"{i+1}. {st_name}")
                with col_del:
                    if st.button("🗑️", key=f"del_st_exp_{i}"):
                        chapter_data['subtopics'].remove(st_name)
                        if st_name in chapter_data['subtopic_data']:
                            del chapter_data['subtopic_data'][st_name]
                        st.rerun()
    
    else:
        is_special_chapter = any(kw in selected_chapter.lower() for kw in ['에필로그', '프롤로그', '서문', '부록', 'epilogue', 'prologue'])
        
        if is_special_chapter:
            st.info(f"📝 '{selected_chapter}'는 소제목 없이 바로 본문을 작성합니다.")
            
            chapter_as_subtopic = selected_chapter
            if chapter_as_subtopic not in chapter_data['subtopic_data']:
                chapter_data['subtopic_data'][chapter_as_subtopic] = {'questions': [], 'answers': [], 'content': ''}
            
            subtopic_data = chapter_data['subtopic_data'][chapter_as_subtopic]
            
            content_widget_key_special = f"content_special_{selected_chapter}"
            
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.markdown('<p class="section-label">Step 01</p>', unsafe_allow_html=True)
                st.markdown(f"### 🎤 인터뷰: {selected_chapter}")
                
                if st.button("🎤 질문 생성하기", key="gen_questions_special"):
                    with st.spinner("질문 생성 중..."):
                        questions_text = generate_interview_questions(
                            selected_chapter, 
                            selected_chapter, 
                            st.session_state['topic']
                        )
                        questions = re.findall(r'Q\d+:\s*(.+)', questions_text)
                        if not questions:
                            questions = [q.strip() for q in questions_text.split('\n') if q.strip() and '?' in q][:3]
                        subtopic_data['questions'] = questions
                        subtopic_data['answers'] = [''] * len(questions)
                        st.rerun()
                
                if subtopic_data['questions']:
                    for i, q in enumerate(subtopic_data['questions']):
                        st.markdown(f"**Q{i+1}.** {q}")
                        if i >= len(subtopic_data['answers']):
                            subtopic_data['answers'].append('')
                        subtopic_data['answers'][i] = st.text_area(
                            f"A{i+1}",
                            value=subtopic_data['answers'][i],
                            key=f"answer_special_{i}",
                            height=80,
                            label_visibility="collapsed"
                        )
                else:
                    st.info("👆 '질문 생성하기' 버튼을 눌러 인터뷰를 시작하세요.")
            
            with col2:
                st.markdown('<p class="section-label">Step 02</p>', unsafe_allow_html=True)
                st.markdown(f"### 📝 본문: {selected_chapter}")
                
                has_answers = subtopic_data.get('questions') and any(a.strip() for a in subtopic_data.get('answers', []))
                
                if has_answers:
                    if st.button("✨ 본문 생성하기", key="gen_content_special"):
                        with st.spinner("집필 중... (30초~1분)"):
                            content = generate_subtopic_content(
                                selected_chapter,
                                selected_chapter,
                                subtopic_data['questions'],
                                subtopic_data['answers'],
                                st.session_state['topic'],
                                st.session_state['target_persona']
                            )
                            st.session_state['chapters'][selected_chapter]['subtopic_data'][chapter_as_subtopic]['content'] = content
                            st.session_state[content_widget_key_special] = content
                            trigger_auto_save()
                            st.rerun()
                else:
                    st.info("👈 먼저 인터뷰 질문에 답변해주세요.")
                
                stored_content = st.session_state['chapters'][selected_chapter]['subtopic_data'].get(chapter_as_subtopic, {}).get('content', '')
                if content_widget_key_special not in st.session_state:
                    st.session_state[content_widget_key_special] = stored_content
                
                edited_content = st.text_area(
                    "본문 내용",
                    height=400,
                    key=content_widget_key_special,
                    label_visibility="collapsed"
                )
                
                if content_widget_key_special in st.session_state:
                    st.session_state['chapters'][selected_chapter]['subtopic_data'][chapter_as_subtopic]['content'] = st.session_state[content_widget_key_special]
                
                final_content = st.session_state['chapters'][selected_chapter]['subtopic_data'].get(chapter_as_subtopic, {}).get('content', '')
                if final_content:
                    char_count = calculate_char_count(final_content)
                    st.caption(f"📊 {char_count:,}자")
                    st.success(f"✅ '{selected_chapter}' 본문 작성 완료!")
        
        else:
            st.warning("⚠️ 이 챕터에 소제목이 없습니다. 아래에서 소제목을 생성하거나 추가해주세요.")
            
            st.markdown("### 📝 소제목 생성")
            
            col_gen, col_add = st.columns(2)
            
            with col_gen:
                st.markdown("**자동 생성**")
                num_subtopics = st.number_input(
                    "생성할 개수",
                    min_value=1,
                    max_value=10,
                    value=3,
                    key="num_subtopics_gen_empty"
                )
                if st.button("✨ 소제목 자동 생성", key="gen_subtopics_empty"):
                    with st.spinner("베스트셀러급 소제목 생성 중..."):
                        subtopics_text = generate_subtopics(
                            selected_chapter,
                            st.session_state['topic'],
                            st.session_state['target_persona'],
                            num_subtopics
                        )
                        new_subtopics = []
                        for line in subtopics_text.split('\n'):
                            line = line.strip()
                            if line and (line[0].isdigit() or line.startswith('-')):
                                cleaned = re.sub(r'^[\d\.\-\s]+', '', line).strip()
                                if cleaned:
                                    new_subtopics.append(cleaned)
                        
                        if new_subtopics:
                            chapter_data['subtopics'] = new_subtopics[:num_subtopics]
                            for st_name in new_subtopics[:num_subtopics]:
                                chapter_data['subtopic_data'][st_name] = {'questions': [], 'answers': [], 'content': ''}
                            st.success(f"✅ {len(new_subtopics[:num_subtopics])}개 소제목 생성됨!")
                            st.rerun()
            
            with col_add:
                st.markdown("**직접 입력**")
                new_subtopic_name = st.text_input(
                    "소제목 이름",
                    placeholder="직접 입력하세요",
                    key="new_subtopic_empty"
                )
                if st.button("➕ 소제목 추가", key="add_subtopic_empty"):
                    if new_subtopic_name.strip():
                        chapter_data['subtopics'].append(new_subtopic_name)
                        chapter_data['subtopic_data'][new_subtopic_name] = {'questions': [], 'answers': [], 'content': ''}
                        st.success(f"'{new_subtopic_name}' 추가됨!")
                        st.rerun()

    st.markdown("---")
    st.markdown("### 📖 작성된 본문 통합 보기")
    
    all_content_display = ""
    content_count_tab4 = 0
    
    for ch_idx, ch in enumerate(st.session_state['outline'], 1):
        if ch in st.session_state['chapters']:
            ch_data = st.session_state['chapters'][ch]
            if 'subtopic_data' in ch_data:
                chapter_has_content = False
                chapter_content_parts = []
                
                subtopic_list = ch_data.get('subtopics', [])
                if not subtopic_list and ch in ch_data['subtopic_data']:
                    subtopic_list = [ch]
                
                for st_name in subtopic_list:
                    st_data = ch_data['subtopic_data'].get(st_name, {})
                    if st_data.get('content'):
                        content_text = st_data['content']
                        chapter_content_parts.append(f"**{st_name}**\n\n{content_text}")
                        content_count_tab4 += 1
                        chapter_has_content = True
                
                if chapter_has_content:
                    all_content_display += f"\n\n---\n\n## {ch}\n\n"
                    all_content_display += "\n\n".join(chapter_content_parts)
    
    pure_content = get_all_content_text()
    
    if pure_content:
        total_chars_tab4 = calculate_char_count(pure_content)
        st.success(f"✅ 총 {content_count_tab4}개 소제목 작성 완료 | {total_chars_tab4:,}자")
        
        with st.expander("📖 전체 본문 펼쳐보기", expanded=False):
            for ch_idx, ch in enumerate(st.session_state['outline'], 1):
                if ch in st.session_state['chapters']:
                    ch_data = st.session_state['chapters'][ch]
                    if 'subtopic_data' in ch_data:
                        subtopic_list = ch_data.get('subtopics', [])
                        if not subtopic_list and ch in ch_data['subtopic_data']:
                            subtopic_list = [ch]
                        
                        chapter_has_content = False
                        chapter_contents = []
                        
                        for st_name in subtopic_list:
                            st_data = ch_data['subtopic_data'].get(st_name, {})
                            if st_data.get('content'):
                                cleaned_content = clean_content_for_display(st_data['content'], st_name, ch)
                                if cleaned_content.strip():
                                    chapter_contents.append((st_name, cleaned_content))
                                    chapter_has_content = True
                        
                        if chapter_has_content:
                            st.markdown(f"## {ch}")
                            st.markdown("---")
                            
                            for st_name, content in chapter_contents:
                                st.markdown(f"**{st_name}**")
                                st.markdown(content)
                                st.markdown("")
    else:
        st.info("💡 아직 작성된 본문이 없습니다. 위에서 소제목을 선택하고 본문을 작성해주세요.")

# === TAB 5: 문체 다듬기 ===
with tabs[4]:
    st.markdown("## 문체 다듬기 & 품질 검사")
    
    has_content = False
    for ch_data in st.session_state['chapters'].values():
        if 'subtopic_data' in ch_data:
            for st_data in ch_data['subtopic_data'].values():
                if st_data.get('content'):
                    has_content = True
                    break
    
    if not has_content:
        st.info("💡 먼저 본문을 작성해주세요. 또는 아래에서 직접 텍스트를 입력할 수 있습니다.")
        
        direct_content = st.text_area(
            "다듬을 텍스트 직접 입력",
            height=300,
            placeholder="다듬고 싶은 텍스트를 여기에 붙여넣으세요..."
        )
        
        if direct_content:
            has_content = True
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown('<p class="section-label">Style</p>', unsafe_allow_html=True)
        st.markdown("### 문체 다듬기")
        
        content_options = []
        for ch in st.session_state['outline']:
            if ch in st.session_state['chapters']:
                ch_data = st.session_state['chapters'][ch]
                if 'subtopic_data' in ch_data:
                    for st_name, st_data in ch_data['subtopic_data'].items():
                        if st_data.get('content'):
                            content_options.append(f"{ch} > {st_name}")
        
        if content_options:
            selected_content = st.selectbox(
                "다듬을 콘텐츠 선택",
                content_options,
                key="refine_select"
            )
        
        style = st.selectbox(
            "목표 스타일",
            ["친근한", "전문적", "직설적", "스토리텔링"],
            key="style_select"
        )
        
        if st.button("✨ 문체 다듬기", key="refine_btn"):
            content_to_refine = ""
            
            if content_options and selected_content:
                parts = selected_content.split(" > ")
                if len(parts) == 2:
                    ch, st_name = parts
                    content_to_refine = st.session_state['chapters'][ch]['subtopic_data'][st_name]['content']
            elif 'direct_content' in dir() and direct_content:
                content_to_refine = direct_content
            
            if content_to_refine:
                with st.spinner("다듬는 중..."):
                    refined = refine_content(content_to_refine, style)
                    st.session_state['refined_content'] = refined
            else:
                st.error("다듬을 콘텐츠를 선택해주세요.")
        
        if st.session_state.get('refined_content'):
            st.text_area("다듬어진 본문", value=st.session_state['refined_content'], height=400)
            
            if st.button("원본에 적용", key="apply_refined"):
                if content_options and selected_content:
                    parts = selected_content.split(" > ")
                    if len(parts) == 2:
                        ch, st_name = parts
                        st.session_state['chapters'][ch]['subtopic_data'][st_name]['content'] = st.session_state['refined_content']
                        trigger_auto_save()
                        st.success("적용됨!")
                        st.rerun()
    
    with col2:
        st.markdown('<p class="section-label">Quality</p>', unsafe_allow_html=True)
        st.markdown("### 품질 검사")
        
        if st.button("🔍 베스트셀러 체크", key="quality_btn"):
            content_to_check = ""
            
            if content_options and selected_content:
                parts = selected_content.split(" > ")
                if len(parts) == 2:
                    ch, st_name = parts
                    content_to_check = st.session_state['chapters'][ch]['subtopic_data'][st_name]['content']
            elif 'direct_content' in dir() and direct_content:
                content_to_check = direct_content
            
            if content_to_check:
                with st.spinner("분석 중..."):
                    quality_result = check_quality(content_to_check)
                    st.session_state['quality_result'] = quality_result
            else:
                st.error("검사할 콘텐츠를 선택해주세요.")
        
        if st.session_state.get('quality_result'):
            st.markdown(f"""
            <div class="info-card">
                {st.session_state['quality_result'].replace(chr(10), '<br>')}
            </div>
            """, unsafe_allow_html=True)

# === TAB 6: 최종 출력 ===
with tabs[5]:
    st.markdown("## 최종 출력 & 마케팅")
    
    col1, col2 = st.columns([1.5, 1])
    
    with col1:
        st.markdown('<p class="section-label">Export</p>', unsafe_allow_html=True)
        st.markdown("### 전자책 다운로드")
        
        book_title = st.text_input("전자책 제목", value=st.session_state.get('book_title', ''), key="final_title")
        subtitle = st.text_input("부제", value=st.session_state.get('subtitle', ''), key="final_subtitle")
        
        st.session_state['book_title'] = book_title
        st.session_state['subtitle'] = subtitle
        
        st.markdown("### 스타일 설정")
        
        col_s1, col_s2 = st.columns(2)
        with col_s1:
            font_family = st.selectbox("본문 폰트", ["Pretendard", "Noto Sans KR", "Nanum Gothic"], key="font_select")
            font_size = st.selectbox("본문 크기", ["16px", "17px", "18px"], key="fontsize_select")
        with col_s2:
            line_height = st.selectbox("줄간격", ["1.8", "1.9", "2.0"], key="lineheight_select")
            max_width = st.selectbox("최대 폭", ["700px", "800px", "900px"], key="maxwidth_select")
        
        with st.expander("상세 설정"):
            title_size = st.selectbox("제목 크기", ["32px", "36px", "40px"], key="titlesize_select")
            chapter_size = st.selectbox("챕터 제목 크기", ["24px", "26px", "28px"], key="chaptersize_select")
            subtopic_size = st.selectbox("소제목 크기", ["18px", "20px", "22px"], key="subtopicsize_select")
            text_color = st.color_picker("본문 색상", "#333333", key="textcolor_select")
        
        st.markdown("---")
        
        full_book_txt = ""
        full_book_html = ""
        
        if book_title:
            full_book_txt += f"{book_title}\n"
            full_book_html += f"<h1>{book_title}</h1>\n"
        if subtitle:
            full_book_txt += f"{subtitle}\n"
            full_book_html += f"<p style='color: #666; font-size: 14px; margin-top: -10px;'>{subtitle}</p>\n"
        
        full_book_txt += "\n" + "="*50 + "\n\n"
        full_book_html += "<hr style='border: none; border-top: 1px solid #ddd; margin: 30px 0;'>\n"
        
        for chapter in st.session_state['outline']:
            if chapter in st.session_state['chapters']:
                ch_data = st.session_state['chapters'][chapter]
                
                if 'subtopic_data' in ch_data:
                    chapter_has_content = False
                    for st_name in ch_data.get('subtopics', []):
                        st_data = ch_data['subtopic_data'].get(st_name, {})
                        if st_data.get('content'):
                            chapter_has_content = True
                            break
                    
                    if chapter_has_content:
                        full_book_txt += f"\n{chapter}\n" + "-"*40 + "\n\n"
                        full_book_html += f"<h2 style='font-size: {chapter_size}; margin-top: 50px;'>{chapter}</h2>\n"
                        
                        for st_name in ch_data.get('subtopics', []):
                            st_data = ch_data['subtopic_data'].get(st_name, {})
                            if st_data.get('content'):
                                full_book_txt += f"\n{st_name}\n\n{st_data['content']}\n\n"
                                
                                paragraphs = st_data['content'].split('\n\n')
                                full_book_html += f"<h3 style='font-size: {subtopic_size}; margin-top: 35px;'>{st_name}</h3>\n"
                                for para in paragraphs:
                                    para = para.strip()
                                    if para:
                                        full_book_html += f"<p style='font-size: {font_size}; line-height: {line_height}; color: {text_color};'>{para}</p>\n"
        
        html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{book_title or '전자책'}</title>
    <link href="https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css" rel="stylesheet">
    <style>
        @page {{
            margin: 2cm;
        }}
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        h1 {{
            font-size: {title_size};
            color: #111;
            margin-bottom: 10px;
            font-weight: 700;
        }}
        h2 {{
            font-size: {chapter_size};
            color: #222;
            margin-top: 50px;
            margin-bottom: 20px;
            font-weight: 700;
        }}
        h3 {{
            font-size: {subtopic_size};
            color: #333;
            margin-top: 35px;
            margin-bottom: 15px;
            font-weight: 700;
        }}
        body {{
            font-family: '{font_family}', sans-serif;
            max-width: {max_width};
            margin: 0 auto;
            padding: 60px 20px;
            word-break: keep-all;
            font-weight: 500;
        }}
    </style>
</head>
<body>
{full_book_html}
</body>
</html>"""
        
        def create_docx():
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                if book_title:
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(book_title)
                    title_run.font.size = Pt(28)
                    title_run.font.bold = True
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if subtitle:
                    sub_para = doc.add_paragraph()
                    sub_run = sub_para.add_run(subtitle)
                    sub_run.font.size = Pt(14)
                    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if book_title or subtitle:
                    doc.add_paragraph()
                
                for chapter in st.session_state['outline']:
                    if chapter in st.session_state['chapters']:
                        ch_data = st.session_state['chapters'][chapter]
                        
                        if 'subtopic_data' in ch_data:
                            chapter_has_content = False
                            for st_name in ch_data.get('subtopics', []):
                                st_data = ch_data['subtopic_data'].get(st_name, {})
                                if st_data.get('content'):
                                    chapter_has_content = True
                                    break
                            
                            if chapter_has_content:
                                ch_para = doc.add_paragraph()
                                ch_run = ch_para.add_run(chapter)
                                ch_run.font.size = Pt(20)
                                ch_run.font.bold = True
                                
                                for st_name in ch_data.get('subtopics', []):
                                    st_data = ch_data['subtopic_data'].get(st_name, {})
                                    if st_data.get('content'):
                                        st_para = doc.add_paragraph()
                                        st_run = st_para.add_run(st_name)
                                        st_run.font.size = Pt(14)
                                        st_run.font.bold = True
                                        
                                        paragraphs = st_data['content'].split('\n\n')
                                        for para in paragraphs:
                                            para = para.strip()
                                            if para:
                                                p = doc.add_paragraph()
                                                run = p.add_run(para)
                                                run.font.size = Pt(11)
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer.getvalue()
            except ImportError:
                return None
        
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            st.download_button(
                "📄 TXT 다운로드",
                full_book_txt,
                file_name=f"{book_title or 'ebook'}_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        with col_dl2:
            st.download_button(
                "🌐 HTML 다운로드",
                html_content,
                file_name=f"{book_title or 'ebook'}_{datetime.now().strftime('%Y%m%d')}.html",
                mime="text/html",
                use_container_width=True
            )
        
        col_dl3, col_dl4 = st.columns(2)
        with col_dl3:
            docx_data = create_docx()
            if docx_data:
                st.download_button(
                    "📘 워드(DOCX) 다운로드",
                    docx_data,
                    file_name=f"{book_title or 'ebook'}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.info("워드 파일: python-docx 필요")
        
        with col_dl4:
            rtf_content = f"""{{\\rtf1\\ansi\\deff0
{{\\fonttbl{{\\f0 맑은 고딕;}}}}
\\f0\\fs24
{book_title}\\par
{subtitle}\\par
\\par
"""
            for chapter in st.session_state['outline']:
                if chapter in st.session_state['chapters']:
                    ch_data = st.session_state['chapters'][chapter]
                    if 'subtopic_data' in ch_data:
                        chapter_has_content = any(ch_data['subtopic_data'].get(st_name, {}).get('content') for st_name in ch_data.get('subtopics', []))
                        if chapter_has_content:
                            rtf_content += f"\\par\\b {chapter}\\b0\\par\\par"
                            for st_name in ch_data.get('subtopics', []):
                                st_data = ch_data['subtopic_data'].get(st_name, {})
                                if st_data.get('content'):
                                    rtf_content += f"\\b {st_name}\\b0\\par"
                                    content = st_data['content'].replace('\n', '\\par ')
                                    rtf_content += f"{content}\\par\\par"
            rtf_content += "}"
            
            st.download_button(
                "📗 RTF 다운로드 (한글호환)",
                rtf_content,
                file_name=f"{book_title or 'ebook'}_{datetime.now().strftime('%Y%m%d')}.rtf",
                mime="application/rtf",
                use_container_width=True
            )
        
        st.caption("💡 RTF 파일은 한글, 워드, 리브레오피스 등에서 열 수 있습니다.")
        
        st.markdown("---")
        
        st.markdown("### 📖 작성된 본문 종합 보기")
        
        all_content = ""
        content_count = 0
        
        for ch_idx, chapter in enumerate(st.session_state['outline'], 1):
            if chapter in st.session_state['chapters']:
                ch_data = st.session_state['chapters'][chapter]
                if 'subtopic_data' in ch_data:
                    chapter_has_content = False
                    chapter_content_parts = []
                    
                    subtopic_list = ch_data.get('subtopics', [])
                    if not subtopic_list and chapter in ch_data['subtopic_data']:
                        subtopic_list = [chapter]
                    
                    for st_name in subtopic_list:
                        st_data = ch_data['subtopic_data'].get(st_name, {})
                        if st_data.get('content'):
                            chapter_content_parts.append(f"**{st_name}**\n\n{st_data['content']}")
                            content_count += 1
                            chapter_has_content = True
                    
                    if chapter_has_content:
                        all_content += f"\n\n---\n\n## {chapter}\n\n"
                        all_content += "\n\n".join(chapter_content_parts)
        
        if all_content:
            st.success(f"✅ 총 {content_count}개 소제목 작성 완료")
            
            pure_content_tab6 = get_all_content_text()
            total_chars = calculate_char_count(pure_content_tab6)
            st.caption(f"📊 총 {total_chars:,}자 / 약 {total_chars//500}페이지 (500자/페이지 기준)")
            
            with st.expander("📖 전체 본문 펼쳐보기", expanded=False):
                for ch_idx, chapter in enumerate(st.session_state['outline'], 1):
                    if chapter in st.session_state['chapters']:
                        ch_data = st.session_state['chapters'][chapter]
                        if 'subtopic_data' in ch_data:
                            subtopic_list = ch_data.get('subtopics', [])
                            if not subtopic_list and chapter in ch_data['subtopic_data']:
                                subtopic_list = [chapter]
                            
                            chapter_has_content = False
                            chapter_contents = []
                            
                            for st_name in subtopic_list:
                                st_data = ch_data['subtopic_data'].get(st_name, {})
                                if st_data.get('content'):
                                    cleaned_content = clean_content_for_display(st_data['content'], st_name, chapter)
                                    if cleaned_content.strip():
                                        chapter_contents.append((st_name, cleaned_content))
                                        chapter_has_content = True
                            
                            if chapter_has_content:
                                st.markdown(f"## {chapter}")
                                st.markdown("---")
                                
                                for st_name, content in chapter_contents:
                                    st.markdown(f"**{st_name}**")
                                    st.markdown(content)
                                    st.markdown("")
            
            with st.expander("✏️ 전체 본문 편집하기 (텍스트)", expanded=False):
                edit_text = ""
                for chapter in st.session_state['outline']:
                    if chapter in st.session_state['chapters']:
                        ch_data = st.session_state['chapters'][chapter]
                        if 'subtopic_data' in ch_data:
                            subtopic_list = ch_data.get('subtopics', [])
                            if not subtopic_list and chapter in ch_data['subtopic_data']:
                                subtopic_list = [chapter]
                            
                            chapter_has_content = False
                            for st_name in subtopic_list:
                                st_data = ch_data['subtopic_data'].get(st_name, {})
                                if st_data.get('content'):
                                    if not chapter_has_content:
                                        edit_text += f"\n\n{'='*50}\n{chapter}\n{'='*50}\n\n"
                                        chapter_has_content = True
                                    cleaned = clean_content_for_display(st_data['content'], st_name, chapter)
                                    edit_text += f"[{st_name}]\n\n{cleaned}\n\n"
                
                edited_all = st.text_area(
                    "전체 본문 (편집 가능)",
                    value=edit_text.strip(),
                    height=600,
                    key="full_content_edit"
                )
                st.caption("여기서 수정한 내용은 개별 소제목에는 반영되지 않습니다. 최종 다운로드용으로만 사용됩니다.")
        else:
            st.info("💡 아직 작성된 본문이 없습니다. '④ 본문 작성' 탭에서 먼저 본문을 작성해주세요.")
        
        st.markdown("---")
        
        if st.button("👁️ 스타일 미리보기", key="preview_btn", use_container_width=True):
            st.session_state['show_preview'] = True
        
        if st.session_state.get('show_preview'):
            st.markdown("### 스타일 미리보기")
            preview_sample = f"""
            <div style="font-family: '{font_family}', sans-serif; max-width: {max_width}; line-height: {line_height}; color: {text_color}; font-size: {font_size}; border: 1px solid #ddd; padding: 30px; border-radius: 8px; background: #fff;">
                <h1 style="font-size: {title_size}; font-weight: 700; color: #111; margin-bottom: 5px;">{book_title or '전자책 제목'}</h1>
                <p style="color: #666; font-size: 14px;">{subtitle or '부제목'}</p>
                <hr style="border: none; border-top: 1px solid #ddd; margin: 20px 0;">
                <h2 style="font-size: {chapter_size}; font-weight: 700; color: #222;">챕터1: 왜 열심히 하는 사람이 가난할까</h2>
                <h3 style="font-size: {subtopic_size}; font-weight: 700; color: #333;">그날 통장 잔고 47만원</h3>
                <p style="font-size: {font_size}; line-height: {line_height};">2019년 3월. 통장 잔고를 확인했습니다. 47만원. 월급날까지 2주. 저는 바닥이었습니다.</p>
                <p style="font-size: {font_size}; line-height: {line_height};">솔직히 말씀드리면, 저도 처음엔 몰랐습니다. 열심히만 하면 되는 줄 알았거든요. 새벽 6시에 일어나서 밤 11시까지 일했습니다. 주말도 없었습니다.</p>
            </div>
            """
            st.markdown(preview_sample, unsafe_allow_html=True)
    
    with col2:
        st.markdown('<p class="section-label">Marketing</p>', unsafe_allow_html=True)
        st.markdown("### 마케팅 카피")
        
        if st.button("카피 생성하기", key="marketing_btn"):
            with st.spinner("생성 중..."):
                marketing = generate_marketing_copy(
                    st.session_state.get('book_title', st.session_state['topic']),
                    st.session_state.get('subtitle', ''),
                    st.session_state['topic'],
                    st.session_state['target_persona']
                )
                st.session_state['marketing_copy'] = marketing
        
        if st.session_state.get('marketing_copy'):
            st.markdown(f"""
            <div class="info-card">
                {st.session_state['marketing_copy'].replace(chr(10), '<br>')}
            </div>
            """, unsafe_allow_html=True)

# --- 자동 저장 처리 ---
if st.session_state.get('auto_save_trigger'):
    st.session_state['auto_save_trigger'] = False
    auto_save_data = get_auto_save_data()
    auto_save_json = json.dumps(auto_save_data, ensure_ascii=False, indent=2)
    file_name = st.session_state.get('book_title', '전자책') or '전자책'
    file_name = re.sub(r'[^\w\s가-힣-]', '', file_name)[:20]
    
    st.toast("💾 자동 저장 준비됨!")
    
    with st.sidebar:
        st.markdown("---")
        st.markdown("### 🔔 자동 저장")
        st.download_button(
            "💾 백업 다운로드",
            auto_save_json,
            file_name=f"자동저장_{file_name}_{datetime.now().strftime('%H%M')}.json",
            mime="application/json",
            use_container_width=True,
            type="primary"
        )
        st.caption("중요 작업 완료됨 - 백업 권장!")

# --- 푸터 ---
st.markdown("""
<div class="premium-footer">
    <span class="premium-footer-text">전자책 작성 프로그램 — </span><span class="premium-footer-author">남현우 작가</span>
</div>
""", unsafe_allow_html=True)
